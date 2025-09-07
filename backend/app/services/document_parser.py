"""
Partnership Agreement Parser Service
Advanced NLP-powered document parsing for partnership agreements
"""
import asyncio
import logging
import uuid
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
import json
import re

# Document processing imports
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument

# AI/ML imports
from openai import AsyncOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import tiktoken

# Custom imports
from app.config import settings
from app.models.document import DocumentParseResult, PartnershipAgreement
from app.utils.text_processing import clean_text, extract_numbers, extract_percentages

logger = logging.getLogger(__name__)

@dataclass
class ExtractionResult:
    """Results from document extraction"""
    confidence: float
    data: Dict
    source_text: str
    page_number: Optional[int] = None

class PartnershipAgreementParser:
    """
    AI-powered parser for partnership agreements
    Extracts key economic terms, allocation rules, and partnership structure
    """
    
    def __init__(self):
        self.openai_client = AsyncOpenAI(api_key=settings.OPENAI_API_KEY)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=4000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""]
        )
        self.encoding = tiktoken.get_encoding("cl100k_base")
        
        # Partnership tax terminology patterns
        self.tax_patterns = {
            "partners": r"(?i)partner(?:s)?|member(?:s)?|equity holder(?:s)?",
            "ownership": r"(?i)ownership percentage|interest|share|units?",
            "capital": r"(?i)capital contribution|capital account|capital call",
            "distribution": r"(?i)distribution|waterfall|preferred return|hurdle",
            "allocation": r"(?i)allocation|special allocation|704\(b\)|substantial economic effect",
            "management": r"(?i)management fee|carry|carried interest|promote",
            "liquidation": r"(?i)liquidation|wind.?up|dissolution"
        }
    
    async def parse_agreement(
        self, 
        content: bytes, 
        filename: str,
        partnership_id: Optional[str] = None
    ) -> DocumentParseResult:
        """
        Main parsing method for partnership agreements
        
        Args:
            content: Raw document bytes
            filename: Original filename
            partnership_id: Associated partnership ID
            
        Returns:
            DocumentParseResult with extracted data and confidence scores
        """
        try:
            document_id = str(uuid.uuid4())
            logger.info(f"Starting parse for document {document_id}: {filename}")
            
            # Extract text from document
            text_content = await self._extract_text(content, filename)
            if not text_content:
                raise ValueError("No text content could be extracted from document")
            
            # Split into manageable chunks
            chunks = self._split_document(text_content)
            
            # Extract structured data using AI
            extraction_results = await self._extract_partnership_data(chunks)
            
            # Validate and consolidate results
            consolidated_data = self._consolidate_extractions(extraction_results)
            
            # Calculate overall confidence score
            confidence_score = self._calculate_confidence(extraction_results)
            
            # Create parse result
            parse_result = DocumentParseResult(
                document_id=document_id,
                partnership_id=partnership_id,
                filename=filename,
                status="completed" if confidence_score > 0.8 else "requires_review",
                confidence_score=confidence_score,
                extracted_data=consolidated_data,
                raw_text=text_content[:5000],  # First 5k chars for reference
                created_at=datetime.utcnow()
            )
            
            logger.info(f"Parse completed for {document_id} with confidence {confidence_score:.2f}")
            return parse_result
            
        except Exception as e:
            logger.error(f"Error parsing document: {str(e)}")
            return DocumentParseResult(
                document_id=str(uuid.uuid4()),
                partnership_id=partnership_id,
                filename=filename,
                status="failed",
                confidence_score=0.0,
                extracted_data={},
                error_message=str(e),
                created_at=datetime.utcnow()
            )
    
    async def _extract_text(self, content: bytes, filename: str) -> str:
        """Extract text content from document bytes"""
        try:
            if filename.lower().endswith('.pdf'):
                return await self._extract_pdf_text(content)
            elif filename.lower().endswith(('.docx', '.doc')):
                return await self._extract_docx_text(content)
            else:
                raise ValueError(f"Unsupported file type: {filename}")
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF using pdfplumber for better formatting"""
        import io
        text_content = ""
        
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\\n--- Page {page_num + 1} ---\\n"
                        text_content += page_text + "\\n"
            
            return clean_text(text_content)
            
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
            # Fallback to PyPDF2
            return await self._extract_pdf_pypdf2(content)
    
    async def _extract_pdf_pypdf2(self, content: bytes) -> str:
        """Fallback PDF extraction using PyPDF2"""
        import io
        text_content = ""
        
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_content += f"\\n--- Page {page_num + 1} ---\\n"
                    text_content += page_text + "\\n"
            
            return clean_text(text_content)
            
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {str(e)}")
            raise
    
    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX document"""
        import io
        try:
            doc = DocxDocument(io.BytesIO(content))
            text_content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\\n"
            
            return clean_text(text_content)
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise
    
    def _split_document(self, text: str) -> List[Document]:
        """Split document into manageable chunks for AI processing"""
        chunks = self.text_splitter.split_text(text)
        return [
            Document(page_content=chunk, metadata={"chunk_id": i})
            for i, chunk in enumerate(chunks)
        ]
    
    async def _extract_partnership_data(self, chunks: List[Document]) -> List[ExtractionResult]:
        """Extract partnership data from document chunks using GPT-4"""
        extraction_results = []
        
        for chunk in chunks:
            try:
                result = await self._analyze_chunk(chunk)
                if result:
                    extraction_results.append(result)
            except Exception as e:
                logger.warning(f"Failed to analyze chunk: {str(e)}")
                continue
        
        return extraction_results
    
    async def _analyze_chunk(self, chunk: Document) -> Optional[ExtractionResult]:
        """Analyze a single document chunk for partnership information"""
        
        system_prompt = \"\"\"You are an expert in partnership tax law and document analysis. 
        Your task is to extract key partnership agreement information from legal text.
        
        Extract the following information if present:
        1. Partnership/Entity name and type
        2. Partner/Member names and ownership percentages
        3. Capital contribution requirements
        4. Distribution waterfall provisions
        5. Management fee structures
        6. Carried interest/promote provisions
        7. Special allocation rules
        8. Liquidation procedures
        9. Section 704(b) related provisions
        10. Tax-related clauses
        
        Return the information as a JSON object with confidence scores (0.0-1.0) for each field.
        If no relevant information is found, return null.
        \"\"\"
        
        user_prompt = f\"\"\"Analyze this partnership agreement text and extract key information:
        
        Text:
        {chunk.page_content}
        
        Return extracted data as JSON with confidence scores, or null if no relevant partnership information is found.
        \"\"\"
        
        try:
            response = await self.openai_client.chat.completions.create(
                model="gpt-4-turbo-preview",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.1,
                max_tokens=2000
            )
            
            content = response.choices[0].message.content.strip()
            
            if content.lower() == "null" or not content:
                return None
            
            # Parse JSON response
            try:
                extracted_data = json.loads(content)
                
                # Calculate average confidence for this chunk
                confidences = []
                for key, value in extracted_data.items():
                    if isinstance(value, dict) and "confidence" in value:
                        confidences.append(value["confidence"])
                
                avg_confidence = sum(confidences) / len(confidences) if confidences else 0.5
                
                return ExtractionResult(
                    confidence=avg_confidence,
                    data=extracted_data,
                    source_text=chunk.page_content
                )
                
            except json.JSONDecodeError:
                logger.warning(f"Failed to parse JSON response: {content}")
                return None
                
        except Exception as e:
            logger.error(f"Error in GPT-4 analysis: {str(e)}")
            return None
    
    def _consolidate_extractions(self, results: List[ExtractionResult]) -> Dict:
        """Consolidate extraction results from multiple chunks"""
        consolidated = {
            "partnership_info": {},
            "partners": [],
            "capital_structure": {},
            "distributions": {},
            "allocations": {},
            "management": {},
            "tax_provisions": {},
            "metadata": {
                "total_chunks_analyzed": len(results),
                "high_confidence_extractions": len([r for r in results if r.confidence > 0.8])
            }
        }
        
        # Combine data from all results
        for result in results:
            for category, data in result.data.items():
                if category in consolidated:
                    if isinstance(consolidated[category], list):
                        consolidated[category].extend(data if isinstance(data, list) else [data])
                    elif isinstance(consolidated[category], dict):
                        consolidated[category].update(data if isinstance(data, dict) else {})
        
        # Remove duplicates and validate data
        consolidated = self._validate_consolidated_data(consolidated)
        
        return consolidated
    
    def _validate_consolidated_data(self, data: Dict) -> Dict:
        """Validate and clean consolidated data"""
        
        # Remove duplicate partners based on name similarity
        if "partners" in data and isinstance(data["partners"], list):
            unique_partners = []
            seen_names = set()
            
            for partner in data["partners"]:
                if isinstance(partner, dict) and "name" in partner:
                    name_clean = re.sub(r'[^a-zA-Z0-9\\s]', '', partner["name"]).lower().strip()
                    if name_clean not in seen_names:
                        unique_partners.append(partner)
                        seen_names.add(name_clean)
            
            data["partners"] = unique_partners
        
        # Validate percentage totals
        if "partners" in data:
            total_ownership = 0
            for partner in data["partners"]:
                if isinstance(partner, dict) and "ownership_percentage" in partner:
                    try:
                        pct = float(partner["ownership_percentage"])
                        total_ownership += pct
                    except (ValueError, TypeError):
                        continue
            
            data["metadata"]["total_ownership_percentage"] = total_ownership
            data["metadata"]["ownership_totals_valid"] = abs(total_ownership - 100) < 5  # 5% tolerance
        
        return data
    
    def _calculate_confidence(self, results: List[ExtractionResult]) -> float:
        """Calculate overall confidence score for the parsing results"""
        if not results:
            return 0.0
        
        # Weight confidence by amount of data extracted
        weighted_confidences = []
        for result in results:
            data_richness = len(str(result.data))  # Simple measure of data richness
            weight = min(data_richness / 1000, 1.0)  # Cap weight at 1.0
            weighted_confidences.append(result.confidence * weight)
        
        return sum(weighted_confidences) / len(weighted_confidences)
    
    async def get_parse_status(self, document_id: str) -> Dict:
        """Get parsing status for a document"""
        # This would typically query a database
        # For now, return a placeholder
        return {
            "document_id": document_id,
            "status": "completed",
            "last_updated": datetime.utcnow().isoformat()
        }
    
    async def validate_and_update(self, document_id: str, validation_data: Dict) -> Dict:
        """Validate and update extracted data"""
        # This would typically update the database with corrected data
        # For now, return the validation data
        return {
            "document_id": document_id,
            "updated_data": validation_data,
            "validation_timestamp": datetime.utcnow().isoformat()
        }